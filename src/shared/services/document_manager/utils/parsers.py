"""File parsing utilities for document import."""
import logging
from typing import List, Tuple, Callable, Optional
import os
from pathlib import Path
import base64
from html.parser import HTMLParser
import io
import html as html_lib

logger = logging.getLogger(__name__)

# Optional imports for file support
try:
    import docx
except ImportError:
    docx = None

try:
    import mammoth
except ImportError:
    mammoth = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None

try:
    from pdf2docx import Converter
except ImportError:
    Converter = None

try:
    import chardet
except ImportError:
    chardet = None


# Guard rails for embedded images pulled from DOCX
MAX_DOCX_IMG_SIZE = 20 * 1024 * 1024  # 20MB cap for base64 embeds
MAX_DOCX_STORE_IMG_SIZE = 50 * 1024 * 1024  # 50MB cap when streaming to store
MAX_DOCX_TOTAL_IMG_BUDGET = 250 * 1024 * 1024  # 250MB aggregate budget per DOCX

# Guard rails for PDF processing
PDF_HIGH_FIDELITY_MAX_BYTES = 50 * 1024 * 1024  # 50MB limit for pdf2docx high-fidelity path
PDF_HTML_TOTAL_CAP = 25 * 1024 * 1024  # 25MB cap for aggregated HTML
PDF_HTML_PER_PAGE_CAP = 2 * 1024 * 1024  # 2MB cap per-page HTML


class HTMLTextExtractor(HTMLParser):
    """Simple HTML parser to extract text with basic layout preservation."""
    def __init__(self):
        """Initialize the extractor with an empty text buffer."""
        super().__init__()
        self.text_parts = []
        
    def handle_data(self, data):
        """Append text data to the buffer."""
        self.text_parts.append(data)
        
    def handle_starttag(self, tag, attrs):
        """Process start tags, adding newlines for breaks."""
        if tag == 'br':
            self.text_parts.append('\n')
            
    def handle_endtag(self, tag):
        """Process end tags, ensuring proper spacing for block elements."""
        if tag in ('p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li', 'tr', 'article', 'section'):
            self.text_parts.append('\n')
            
    def get_text(self):
        """Return the aggregated, suppressed text content."""
        return "".join(self.text_parts).strip()

class DocumentParser:
    """Parses various file formats to extract text and HTML."""
    
    @staticmethod
    def parse_file(
        file_path: str,
        store_image_callback: Optional[Callable[[bytes, str], int]] = None,
        high_fidelity: bool = False,
    ) -> tuple[str, str, str, dict]:
        """
        Parse a file and return (content_text, raw_html, file_type, metadata).
        
        Args:
            file_path: Path to the file.
            
        Returns:
            Tuple of (extracted_text, raw_html, file_extension, metadata_dict)
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        metadata = {}
        
        if ext == '.txt':
            text = DocumentParser._parse_txt(path)
            escaped = html_lib.escape(text, quote=True)
            metadata["pipeline"] = "txt:plain"
            metadata["image_mode"] = "base64"
            return text, f"<pre>{escaped}</pre>", 'txt', metadata
        elif ext == '.html' or ext == '.htm':
            html = DocumentParser._parse_html(path)
            
            # Extract text properly
            extractor = HTMLTextExtractor()
            extractor.feed(html)
            text = extractor.get_text()
            
            metadata["pipeline"] = "html:raw"
            metadata["image_mode"] = "base64"
            return text, html, 'html', metadata
        elif ext == '.docx':
            return DocumentParser._parse_docx(path, store_image_callback=store_image_callback)
        elif ext == '.pdf':
            text, html, file_type, metadata = DocumentParser._parse_pdf(path, high_fidelity=high_fidelity)
            # PDF parsers currently embed images as base64 in HTML output
            metadata["image_mode"] = metadata.get("image_mode", "base64")
            return text, html, file_type, metadata
        elif ext == '.rtf':
            text = DocumentParser._parse_rtf(path)
            escaped = html_lib.escape(text, quote=True)
            metadata["pipeline"] = "rtf:striprtf"
            metadata["image_mode"] = "base64"
            return text, f"<pre>{escaped}</pre>", 'rtf', metadata
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    @staticmethod
    def _parse_txt(path: Path) -> str:
        data = path.read_bytes()

        # Prefer detected encoding when available; fall back through utf-8 then latin-1 before ignoring errors.
        encoding = 'utf-8'
        if chardet:
            try:
                detected = chardet.detect(data[:65536])
                if detected and detected.get('encoding'):
                    encoding = detected['encoding']
            except Exception:
                pass

        for enc in (encoding, 'utf-8', 'latin-1'):
            try:
                return data.decode(enc)
            except Exception:
                continue

        return data.decode('utf-8', errors='ignore')

    @staticmethod
    def _parse_html(path: Path) -> str:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()


    @staticmethod
    def _get_run_images_as_html(
        run_element,
        parts_proxy,
        store_image_callback: Optional[Callable[[bytes, str], int]] = None,
        image_budget: Optional[dict] = None,
        warnings: Optional[list] = None,
    ) -> list[str]:
        """Extract images from a run element and return img tags.

        If store_image_callback is provided, images are stored immediately and emitted
        as docimg:// references; otherwise they are embedded as base64.
        """
        image_htmls = []
        for elem in run_element.iter():
            if elem.tag.endswith('blip'):
                embed_attr = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if embed_attr:
                    try:
                        image_part = parts_proxy.related_parts[embed_attr]
                        image_data = image_part.blob
                        # Size policy: allow larger assets when streaming to store, cap base64 embeds
                        if store_image_callback:
                            if len(image_data) > MAX_DOCX_STORE_IMG_SIZE:
                                logger.warning("Embedded DOCX image exceeds store cap; skipping image")
                                if warnings is not None:
                                    warnings.append("docx_image_store_cap_exceeded")
                                continue
                        else:
                            if len(image_data) > MAX_DOCX_IMG_SIZE:
                                logger.warning("Embedded DOCX image exceeds base64 cap; skipping embed")
                                image_htmls.append("<span class='docimg-omitted' aria-label='image omitted: too large'></span>")
                                if warnings is not None:
                                    warnings.append("docx_image_embed_cap_exceeded")
                                continue
                        if image_budget is not None:
                            used = image_budget.get("used", 0)
                            limit = image_budget.get("limit", MAX_DOCX_TOTAL_IMG_BUDGET)
                            if used + len(image_data) > limit:
                                logger.warning("Embedded DOCX images exceed aggregate budget; skipping image")
                                image_htmls.append("<span class='docimg-omitted' aria-label='image omitted: budget exceeded'></span>")
                                if warnings is not None:
                                    warnings.append("docx_image_budget_exceeded")
                                continue
                        header = image_data[:12]
                        mime = None
                        if header.startswith(b'\xff\xd8'):
                            mime = 'image/jpeg'
                        elif header.startswith(b'\x89PNG'):
                            mime = 'image/png'
                        elif header.startswith(b'GIF8'):
                            mime = 'image/gif'
                        elif header.startswith(b'RIFF') and header[8:12] == b'WEBP':
                            mime = 'image/webp'
                        elif header.startswith(b'BM'):
                            mime = 'image/bmp'

                        if not mime:
                            logger.warning("Embedded DOCX image has unsupported/unknown format; skipping")
                            image_htmls.append("<span class='docimg-omitted' aria-label='image omitted: unsupported format'></span>")
                            if warnings is not None:
                                warnings.append("docx_image_unsupported_format")
                            continue
                        if image_budget is not None:
                            image_budget["used"] = image_budget.get("used", 0) + len(image_data)
                        if store_image_callback:
                            try:
                                image_id = store_image_callback(image_data, mime)
                                image_htmls.append(f'<img src="docimg://{image_id}" />')
                            except Exception:
                                logger.debug("Failed to store image from DOCX run", exc_info=True)
                                # If storage fails, fall back to base64 embed within the safe cap
                                if len(image_data) <= MAX_DOCX_IMG_SIZE:
                                    b64 = base64.b64encode(image_data).decode('utf-8')
                                    image_htmls.append(f'<img src="data:{mime};base64,{b64}" />')
                                else:
                                    image_htmls.append("<span class='docimg-omitted' aria-label='image omitted: too large'></span>")
                        else:
                            b64 = base64.b64encode(image_data).decode('utf-8')
                            image_htmls.append(f'<img src="data:{mime};base64,{b64}" />')
                    except Exception as e:
                        logger.debug("Failed to extract image from DOCX run", exc_info=True)
        return image_htmls

    @staticmethod
    def _process_docx_paragraph(para, tag_override=None, store_image_callback: Optional[Callable[[bytes, str], int]] = None, image_budget: Optional[dict] = None, warnings: Optional[list] = None) -> tuple[str, bool]:
        """
        Convert a docx Paragraph to HTML string.
        Returns: (html_string, is_empty_boolean)
        """
        style_name = para.style.name.lower()
        tag = tag_override if tag_override else 'p'
        
        if not tag_override:
            if 'heading 1' in style_name: tag = 'h1'
            elif 'heading 2' in style_name: tag = 'h2'
            elif 'heading 3' in style_name: tag = 'h3'
            elif 'heading 4' in style_name: tag = 'h4'
            elif 'title' in style_name: tag = 'h1'
            elif 'code' in style_name: tag = 'pre'
        
        # Alignment
        align_attr = "left"
        if para.alignment:
            if 'CENTER' in str(para.alignment): align_attr = "center"
            elif 'JUSTIFY' in str(para.alignment): align_attr = "justify"
            elif 'RIGHT' in str(para.alignment): align_attr = "right"

        para_content = []
        has_content = False
        
        for run in para.runs:
            run_text = run.text
            
            # Handle Images
            images = DocumentParser._get_run_images_as_html(run.element, run.part, store_image_callback=store_image_callback, image_budget=image_budget, warnings=warnings)
            if images:
                para_content.extend(images)
                has_content = True
            
            if not run_text:
                continue
                
            run_text = html_lib.escape(run_text, quote=True).replace('\t', '&nbsp;&nbsp;&nbsp;&nbsp;').replace('\ufffc', '')
            
            # If run has text (even spaces, though we might want to trim? for now keep strict)
            if run_text.strip():
                has_content = True

            styles = []
            if run.bold: styles.append("font-weight: bold")
            if run.italic: styles.append("font-style: italic")
            if run.underline: styles.append("text-decoration: underline")
            
            # Fonts
            font_name = run.font.name
            if not font_name and para.style and para.style.font:
                font_name = para.style.font.name
            
            if font_name:
                styles.append(f"font-family: '{font_name}'")

            # Colors
            if run.font.color and run.font.color.rgb:
                styles.append(f"color: #{run.font.color.rgb}")
            
            if styles:
                style_str = "; ".join(styles)
                para_content.append(f'<span style="{style_str}">{run_text}</span>')
            else:
                para_content.append(run_text)
        
        # Paragraph Formatting (Margins/Line Height)
        # pdf2docx can sometimes produce massive space_after values (~1.5 million EMUs)
        # 12700 EMUs = 1 pt. 1581150 EMUs approx 125pt or 1.7 inches.
        # We will clamp these values to prevent huge gaps.
        
        MAX_MARGIN_PT = 24  # Cap margin-bottom at 24pt
        MAX_LINE_HEIGHT = 2.0 # Cap relative line-height at 2.0
        
        para_style_parts = []
        if para.paragraph_format:
            fmt = para.paragraph_format
            
            # Space After (Margin Bottom)
            if fmt.space_after is not None:
                # Value is usually in EMUs or Pt. Python-docx Length object behaves like int (EMU).
                # 914400 EMUs per inch. 12700 EMUs per point.
                # If value is huge (> 50pt), likely an artifact.
                try:
                    pt_value = fmt.space_after.pt
                    if pt_value > MAX_MARGIN_PT:
                        # print(f"DEBUG: Clamping huge paragraph spacing: {pt_value}pt -> {MAX_MARGIN_PT}pt")
                        pt_value = MAX_MARGIN_PT
                    
                    if pt_value > 0:
                        para_style_parts.append(f"margin-bottom: {pt_value}pt")
                except Exception:
                    pass

            # Space Before (Margin Top)
            if fmt.space_before is not None:
                try:
                    pt_value = fmt.space_before.pt
                    if pt_value > MAX_MARGIN_PT:
                        pt_value = MAX_MARGIN_PT
                    
                    if pt_value > 0:
                        para_style_parts.append(f"margin-top: {pt_value}pt")
                except Exception:
                    pass
            
            # Line Spacing (Line Height)
            if fmt.line_spacing is not None:
                # If float, it's relative (e.g. 1.5). If int, it's exact (Twips or similar? wait docx uses lengths)
                # python-docx: line_spacing property: "If a float, it represents the number of lines... If a Length, it represents the specific height."
                if isinstance(fmt.line_spacing, float):
                     val = fmt.line_spacing
                     if val > MAX_LINE_HEIGHT: val = MAX_LINE_HEIGHT
                     para_style_parts.append(f"line-height: {val}")
                # Ignoring exact line spacing for now to avoid complexity, letting browser handle default unless relative.

        para_style_attr = ""
        if para_style_parts:
            para_style_attr = f" style='{'; '.join(para_style_parts)}'"

        if para_content:
            return f"<{tag} dir='ltr' align='{align_attr}'{para_style_attr}>{''.join(para_content)}</{tag}>", not has_content
        else:
            return f"<{tag} dir='ltr' align='{align_attr}'{para_style_attr}><br/></{tag}>", True

    @staticmethod
    def _process_docx_table(table, store_image_callback: Optional[Callable[[bytes, str], int]] = None, image_budget: Optional[dict] = None, warnings: Optional[list] = None) -> str:
        """Convert a docx Table to HTML string with shading and font support."""
        rows_html = []
        for row in table.rows:
            cells_html = []
            for cell in row.cells:
                # Recursively process cell content
                cell_inner_html = ""
                
                # Use iter_inner_content to handle mixed content (paragraphs + nested tables)
                # Apply whitespace collapsing logic (same as main doc loop)
                empty_para_count = 0
                MAX_EMPTY_PARAS = 1
                
                try:
                    # iter_inner_content() is preferred if available
                    content_iter = cell.iter_inner_content()
                except AttributeError:
                    # Fallback for older python-docx versions
                    content_iter = cell.paragraphs
                
                for block in content_iter:
                    if isinstance(block, docx.text.paragraph.Paragraph):
                        html, is_empty = DocumentParser._process_docx_paragraph(block, store_image_callback=store_image_callback, image_budget=image_budget, warnings=warnings)
                        if is_empty:
                            empty_para_count += 1
                            if empty_para_count <= MAX_EMPTY_PARAS:
                                cell_inner_html += html
                        else:
                            empty_para_count = 0
                            cell_inner_html += html
                    elif isinstance(block, docx.table.Table):
                        # Nested table
                        empty_para_count = 0
                        cell_inner_html += DocumentParser._process_docx_table(block, store_image_callback=store_image_callback, image_budget=image_budget, warnings=warnings)
                
                # Extract Cell Properties (Shading/Background)
                style_attr = "border: 1px solid #ccc; padding: 5px; vertical-align: top;"
                tc = cell._tc
                if tc is not None:
                    tcPr = tc.get_or_add_tcPr()
                    shd = tcPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
                    if shd is not None:
                        fill = shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill")
                        if fill and fill != 'auto':
                            style_attr += f" background-color: #{fill};"
                
                cells_html.append(f"<td style='{style_attr}'>{cell_inner_html}</td>")
            rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
        return f"<table dir='ltr' style='border-collapse: collapse; width: 100%; border: 1px solid #ddd; margin: 10px 0;'>{''.join(rows_html)}</table>"

    @staticmethod
    def _build_numbering_map(doc) -> dict:
        """
        Builds a map of (numId, ilvl) -> 'ul' | 'ol' identifying list types by numbering definition.
        """
        num_map = {} # (numId, ilvl) -> list_type
        
        try:
            # Safely access numbering part
            part = doc.part
            numbering_part = getattr(part, 'numbering_part', None)
            
            if numbering_part is None: 
                return {}
                
            # Access the CT_Numbering element (<w:numbering>)
            element = numbering_part.element
            
            # 1. Map abstractNumId -> { ilvl: format_val }
            abstract_map = {}
            
            # access abstractNum_lst (python-docx generated property)
            for abstract_num in getattr(element, 'abstractNum_lst', []):
                 try:
                     abst_id = abstract_num.abstractNumId.val
                     abstract_map[abst_id] = {}
                     
                     for lvl in abstract_num.lvl_lst:
                         ilvl = int(lvl.ilvl)
                         if lvl.numFmt:
                             abstract_map[abst_id][ilvl] = lvl.numFmt.val
                 except Exception:
                     continue
            
            # 2. Link numId -> abstractNumId -> format
            for num in getattr(element, 'num_lst', []):
                try:
                    num_id = int(num.numId.val)
                    abst_id = num.abstractNumId.val
                    
                    if abst_id in abstract_map:
                        for ilvl, fmt in abstract_map[abst_id].items():
                            # Determine type: bullet -> ul, anything else -> ol
                            # Common formats: decimal, upperRoman, lowerLetter, bullet
                            if fmt == 'bullet':
                                num_map[(num_id, ilvl)] = 'ul'
                            else:
                                num_map[(num_id, ilvl)] = 'ol'
                except Exception:
                    continue
                            
        except Exception as e:
            # print(f"DEBUG: Failed to build numbering map: {e}")
            pass 
            
        return num_map

    @staticmethod
    def _get_list_properties(para, numbering_map=None) -> tuple[str, int]:
        """
        Determine if paragraph is a list item.
        Returns: (list_type ('ul'|'ol'|None), level (0-indexed))
        """
        try:
            # Check for direct numbering properties (numPr)
            if para._p.pPr.numPr is not None:
                # Get indent level
                level = 0
                if para._p.pPr.numPr.ilvl is not None:
                    level = int(para._p.pPr.numPr.ilvl.val)
                
                # High-Precision Check: Use the map if available
                if numbering_map and para._p.pPr.numPr.numId is not None:
                    try:
                        num_id = int(para._p.pPr.numPr.numId.val)
                        if (num_id, level) in numbering_map:
                            return numbering_map[(num_id, level)], level
                    except Exception:
                        pass
                
                # Fallback: Heuristic
                style_name = para.style.name.lower()
                list_type = 'ul'
                if 'num' in style_name or 'order' in style_name:
                    list_type = 'ol'
                
                return list_type, level
            
            # Fallback based on style name
            style_name = para.style.name.lower()
            if 'list paragraph' in style_name or 'bullet' in style_name:
                 return 'ul', 0
                 
        except Exception:
            pass
            
        return None, 0

    @staticmethod
    def _extract_full_text(doc) -> str:
        """
        Extract all text from a DOCX document in reading order,
        including paragraphs and table cells (with nested table support).
        """
        text_parts = []
        
        def extract_table_text(table) -> list[str]:
            """Recursively extract text from a table, including nested tables."""
            table_texts = []
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_texts = []
                    try:
                        # Prefer iter_inner_content for mixed content
                        content_iter = cell.iter_inner_content()
                    except AttributeError:
                        # Fallback for older python-docx
                        content_iter = cell.paragraphs
                    
                    for block in content_iter:
                        if isinstance(block, docx.text.paragraph.Paragraph):
                            if block.text.strip():
                                cell_texts.append(block.text.strip())
                        elif isinstance(block, docx.table.Table):
                            # Nested table - recurse
                            nested = extract_table_text(block)
                            cell_texts.extend(nested)
                    
                    if cell_texts:
                        row_texts.append(" ".join(cell_texts))
                
                if row_texts:
                    # Join cells with pipe structure for visibility (Markdown-ish)
                    table_texts.append(" | ".join(row_texts))
            
            return table_texts
        
        # Iterate document in order (fallback to paragraphs/tables list if iter_inner_content missing)
        try:
            content_iter = doc.iter_inner_content()
        except AttributeError:
            # Fallback: concatenate paragraphs then tables in document order best-effort
            content_iter = list(doc.paragraphs) + list(doc.tables)

        for block in content_iter:
            if isinstance(block, docx.text.paragraph.Paragraph):
                if block.text.strip():
                    text_parts.append(block.text.strip())
            elif isinstance(block, docx.table.Table):
                table_lines = extract_table_text(block)
                text_parts.extend(table_lines)
        
        return "\n".join(text_parts)

    @staticmethod
    def _parse_docx(
        source: object,
        store_image_callback: Optional[Callable[[bytes, str], int]] = None,
    ) -> tuple[str, str, str, dict]:
        """
        Parse a DOCX file from a path or file-like object.
        Args:
            source: Path (str/Path) or file-like object (BytesIO).
        """
        metadata = {}
        text = ""
        html = ""

        # Optimized imports
        if not docx:
            if not mammoth:
                raise ImportError("python-docx or mammoth is required for .docx files")
            # Fallback to pure mammoth if python-docx is missing
            if hasattr(source, 'read'):
                docx_file = source
            else:
                docx_file = open(source, "rb")
                
            try:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
                if hasattr(docx_file, 'seek'):
                    docx_file.seek(0)
                text_result = mammoth.extract_raw_text(docx_file)
                text = text_result.value
                metadata["pipeline"] = "docx:mammoth-only"
                metadata["image_mode"] = "base64"
            finally:
                if not hasattr(source, 'read'):
                    docx_file.close()
                    
            return text, html, 'docx', metadata

        try:
            # python-docx accepts path (str) or file-like object
            doc_source = str(source) if isinstance(source, (str, Path)) else source
            doc = docx.Document(doc_source)
            
            # Extract Metadata
            core_props = doc.core_properties
            if core_props.title: metadata['title'] = core_props.title
            if core_props.author: metadata['author'] = core_props.author
            if core_props.created: metadata['created'] = core_props.created
            
            # --- Text Extraction ---
            text = DocumentParser._extract_full_text(doc)

            warnings: list[str] = []
            image_budget = {"used": 0, "limit": MAX_DOCX_TOTAL_IMG_BUDGET}
            
            # --- Custom HTML Generation (Preserving Fonts & Order) ---
            html_parts = []
            
            # Optimization: Collapse consecutive empty paragraphs
            empty_para_count = 0
            MAX_EMPTY_PARAS = 1 # Allow at most 1 empty line for spacing
            
            # List State Tracking
            list_stack = [] 
            
            # Build precise numbering map
            numbering_map = DocumentParser._build_numbering_map(doc)
            
            # Iterate through document content in order (version-safe)
            try:
                content_iter = doc.iter_inner_content()
            except AttributeError:
                content_iter = list(doc.paragraphs) + list(doc.tables)

            for block in content_iter:
                # --- LIST STATE MANAGEMENT ---
                is_list_item = False
                list_type, list_level = None, 0
                
                if isinstance(block, docx.text.paragraph.Paragraph):
                   list_type, list_level = DocumentParser._get_list_properties(block, numbering_map)
                
                if list_type:
                    is_list_item = True
                    # Reset empty paragraphs if we hit a list (lists usually dense)
                    empty_para_count = 0 
                    
                    # 1. Close lists if we dropped levels
                    while len(list_stack) > (list_level + 1):
                        closing_type = list_stack.pop()
                        html_parts.append(f"</{closing_type}>")
                    
                    # 2. Open lists if we increased levels
                    # Ensure we have parent lists for intermediate levels if needed (simplify to just opening needed levels)
                    while len(list_stack) < (list_level + 1):
                         list_stack.append(list_type)
                         html_parts.append(f"<{list_type} dir='ltr'>")
                    
                    # 3. Handle type switch at same level (uncommon but possible)
                    if list_stack[-1] != list_type:
                        old_type = list_stack.pop()
                        html_parts.append(f"</{old_type}>")
                        list_stack.append(list_type)
                        html_parts.append(f"<{list_type} dir='ltr'>")
                        
                else:
                    # Not a list item: Close ALL open lists
                    while list_stack:
                        closing_type = list_stack.pop()
                        html_parts.append(f"</{closing_type}>")

                # --- BLOCK PROCESSING ---
                if isinstance(block, docx.text.paragraph.Paragraph):
                    # Pass 'li' if it's a list item
                    tag_override = 'li' if is_list_item else None
                    html, is_empty = DocumentParser._process_docx_paragraph(block, tag_override=tag_override, store_image_callback=store_image_callback, image_budget=image_budget, warnings=warnings)
                    
                    if is_empty:
                        empty_para_count += 1
                        if empty_para_count <= MAX_EMPTY_PARAS:
                            html_parts.append(html)
                    else:
                        empty_para_count = 0
                        html_parts.append(html)
                        
                elif isinstance(block, docx.table.Table):
                    # Reset empty count on table
                    empty_para_count = 0
                    html_parts.append(DocumentParser._process_docx_table(block, store_image_callback=store_image_callback, image_budget=image_budget, warnings=warnings))
            
            # Final cleanup: Close any remaining lists
            while list_stack:
                closing_type = list_stack.pop()
                html_parts.append(f"</{closing_type}>")

            html = "\n".join(html_parts)
            metadata["pipeline"] = metadata.get("pipeline", "docx:python-docx")
            metadata["image_mode"] = "docimg" if store_image_callback else "base64"
            if warnings:
                metadata["warnings"] = warnings
            
        except Exception as e:
            logger.exception("Custom DOCX parsing failed; falling back")
            if mammoth:
                if hasattr(source, 'read'):
                    source.seek(0)
                    docx_file = source
                else:
                    docx_file = open(source, "rb")
                    
                try:
                    result = mammoth.convert_to_html(docx_file)
                    html = result.value
                    if hasattr(docx_file, 'seek'):
                        docx_file.seek(0)
                    text_result = mammoth.extract_raw_text(docx_file)
                    text = text_result.value
                    metadata["pipeline"] = "docx:mammoth-fallback"
                    metadata["image_mode"] = "base64"
                finally:
                    if not hasattr(source, 'read'):
                        docx_file.close()
            else:
                escaped = html_lib.escape(text, quote=True)
                html = f"<pre>{escaped}</pre>"
                metadata["pipeline"] = "docx:pre-fallback"
                metadata["image_mode"] = "base64"

        return text, html, 'docx', metadata

    @staticmethod
    def _parse_pdf(path: Path, high_fidelity: bool = False) -> tuple[str, str, str, dict]:
        text_parts: list[str] = []
        html_parts: list[str] = []
        metadata = {}
        warnings: list[str] = []
        
        # Try pdf2docx -> mammoth pipeline for better table/layout preservation
        if Converter and mammoth and docx and high_fidelity:
            try:
                # SAFEGUARD: Skip pdf2docx for large files (>10MB) to prevent hanging/OOM
                file_size = os.path.getsize(path)
                if file_size > PDF_HIGH_FIDELITY_MAX_BYTES: # 10MB
                    logger.info(
                        "PDF > 10MB (%s bytes), skipping high-fidelity conversion. Using fallback.",
                        file_size,
                    )
                    raise ValueError("File too large for expensive conversion")

                # OPTIMIZATION: Use in-memory ByteIO stream instead of temp file
                # avoiding disk I/O improves performance and kept cleaner filesystem
                docx_stream = io.BytesIO()
                
                # Convert PDF to DOCX stream
                # multi_processing=False required for stream operations
                cv = Converter(str(path))
                try:
                    cv.convert(docx_stream, start=0, multi_processing=False)
                finally:
                    cv.close()
                
                # Parse the generated DOCX from memory
                docx_stream.seek(0)
                text, html, _, _ = DocumentParser._parse_docx(docx_stream)
                
                # Now extract real PDF metadata using pypdf or fitz
                metadata = DocumentParser._extract_pdf_metadata(path)
                
                metadata["pipeline"] = "pdf:pdf2docx+mammoth"
                return text, html, 'pdf', metadata
                
            except Exception as e:
                logger.exception("pdf2docx stream conversion failed; falling back")
                # Fall through to fitz/pypdf fallback

        # Use PyMuPDF (fitz) for HTML with images
        if fitz:
            doc = None
            try:
                doc = fitz.open(path)
                
                # Metadata
                meta = doc.metadata
                if meta:
                    if meta.get('title'): metadata['title'] = meta['title']
                    if meta.get('author'): metadata['author'] = meta['author']
                
                html_budget = 0
                for page in doc:
                    # Get text
                    text_parts.append(str(page.get_text("text")))
                    # Get HTML
                    page_html_raw = str(page.get_text("html"))
                    if len(page_html_raw) > PDF_HTML_PER_PAGE_CAP:
                        warnings.append("pdf_page_html_truncated")
                        page_html_raw = page_html_raw[:PDF_HTML_PER_PAGE_CAP]

                    page_len = len(page_html_raw)
                    remaining = PDF_HTML_TOTAL_CAP - html_budget
                    if remaining <= 0:
                        warnings.append("pdf_total_html_truncated")
                        break

                    if page_len > remaining:
                        warnings.append("pdf_total_html_truncated")
                        page_html_raw = page_html_raw[:remaining]
                        page_len = len(page_html_raw)
                    html_budget += page_len
                    # Wrap page content in LTR div to enforce directionality even in fallback
                    if page_html_raw:
                        html_parts.append(f"<div dir='ltr' align='left'>{page_html_raw}</div>")
                    
                text = "\n".join(text_parts)
                html = "\n".join(html_parts)
                metadata["pipeline"] = "pdf:fitz"
                if warnings:
                    metadata["warnings"] = warnings
            except Exception as e:
                # Handle encrypted pdfs etc
                logger.exception("PyMuPDF failed")
                raise ValueError(f"Failed to parse PDF: {e}")
            finally:
                if doc:
                    doc.close()
                
        elif pypdf:
            # Fallback to text only
            try:
                with open(path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    
                    if reader.is_encrypted:
                        try:
                            reader.decrypt('')
                        except Exception:
                            raise ValueError("PDF is encrypted and cannot be read.")

                    # Metadata
                    if reader.metadata:
                        if reader.metadata.title: metadata['title'] = reader.metadata.title
                        if reader.metadata.author: metadata['author'] = reader.metadata.author

                    for page in reader.pages:
                        text_parts.append(page.extract_text() or "")
            except Exception as e:
                raise ValueError(f"pypdf failed: {e}")
            
            text = "\n".join(text_parts)
            escaped = html_lib.escape(text, quote=True)
            html = f"<pre dir='ltr' align='left'>{escaped}</pre>"
            metadata["pipeline"] = "pdf:pypdf"
        else:
            raise ImportError("PyMuPDF or pypdf is required for .pdf files")

        return text, html, 'pdf', metadata

    @staticmethod
    def _extract_pdf_metadata(path: Path) -> dict:
        metadata = {}
        if fitz:
            doc = None
            try:
                doc = fitz.open(path)
                meta = doc.metadata
                if meta:
                    if meta.get('title'): metadata['title'] = meta['title']
                    if meta.get('author'): metadata['author'] = meta['author']
            except Exception:
                pass
            finally:
                if doc:
                    doc.close()
        elif pypdf:
             try:
                with open(path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    if reader.metadata:
                        if reader.metadata.title: metadata['title'] = reader.metadata.title
                        if reader.metadata.author: metadata['author'] = reader.metadata.author
             except Exception:
                 pass
        return metadata

    @staticmethod
    def _parse_rtf(path: Path) -> str:
        if not rtf_to_text:
            raise ImportError("striprtf is required for .rtf files")
            
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            return rtf_to_text(content)