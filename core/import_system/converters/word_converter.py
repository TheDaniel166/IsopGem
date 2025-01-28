from docx import Document
import os
from PyQt5.QtGui import QImage

class WordConverter:
    def convert_to_rtf(self, file_path):
        """Convert Word document to RTF format with enhanced formatting"""
        try:
            doc = Document(file_path)
            rtf_content = []
            image_counter = 0
            
            # RTF Header with font table and color table
            rtf_header = (
                "{\\rtf1\\ansi\\deff0 "
                "{\\fonttbl{\\f0\\froman Times New Roman;}{\\f1\\fswiss Arial;}{\\f2\\fmodern Courier New;}} "
                "{\\colortbl;\\red0\\green0\\blue0;\\red255\\green0\\blue0;\\red0\\green0\\blue255;} "
            )
            rtf_content.append(rtf_header)
            
            for paragraph in doc.paragraphs:
                # Start paragraph
                para_format = []
                
                # Paragraph alignment
                if paragraph.alignment == 1:  # center
                    para_format.append("\\qc")
                elif paragraph.alignment == 2:  # right
                    para_format.append("\\qr")
                elif paragraph.alignment == 3:  # justify
                    para_format.append("\\qj")
                
                # Handle heading styles
                if paragraph.style.name.startswith('Heading'):
                    level = paragraph.style.name[-1]
                    para_format.append(f"\\s{level}\\outlinelvl{level}")
                
                # Start paragraph with formatting
                rtf_content.append("\\n{" + ''.join(para_format))
                
                # Process runs (text with consistent formatting)
                for run in paragraph.runs:
                    run_format = []
                    
                    # Character formatting
                    if run.bold:
                        run_format.append("\\b")
                    if run.italic:
                        run_format.append("\\i")
                    if run.underline:
                        run_format.append("\\ul")
                    if run.font.strike:
                        run_format.append("\\strike")
                    if run.font.superscript:
                        run_format.append("\\super")
                    if run.font.subscript:
                        run_format.append("\\sub")
                    
                    # Font size (convert to half-points)
                    if run.font.size:
                        size = int(run.font.size.pt * 2)
                        run_format.append(f"\\fs{size}")
                    
                    # Font name
                    if run.font.name:
                        run_format.append("\\f1")  # Using Arial from font table
                    
                    # Handle images
                    if run.element.findall('.//w:drawing') or run.element.findall('.//w:pict'):
                        try:
                            image_rel = next(
                                rel for rel in paragraph.part.rels.values() 
                                if rel.reltype.endswith('/image')
                            )
                            image_path = os.path.join(
                                os.path.dirname(file_path),
                                image_rel.target
                            )
                            if os.path.exists(image_path):
                                image = QImage(image_path)
                                if not image.isNull():
                                    image_counter += 1
                                    rtf_content.append(
                                        "{\\pict\\jpegblip"
                                        f"\\picw{image.width()}\\pich{image.height()}"
                                        f"\\picwgoal{image.width()*20}\\pichgoal{image.height()*20}"
                                        "}"
                                    )
                        except (StopIteration, Exception):
                            pass
                    
                    # Add text with formatting
                    text = run.text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                    if run_format:
                        rtf_content.append("{" + ''.join(run_format) + " " + text + "}")
                    else:
                        rtf_content.append(text)
                
                # End paragraph
                rtf_content.append("\\par}")
            
            # Handle tables
            for table in doc.tables:
                rtf_content.append("\\par{")  # Start table
                for row in table.rows:
                    rtf_content.append("\\trowd\\trgaph144")  # Start row
                    for cell in row.cells:
                        rtf_content.append("\\cellx2000")  # Simple fixed width cells
                        rtf_content.append("\\intbl ")
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                            rtf_content.append(f"{text}\\cell ")
                    rtf_content.append("\\row")  # End row
                rtf_content.append("}")  # End table
            
            # Close RTF document
            rtf_content.append("}")
            
            return ''.join(rtf_content)
            
        except Exception as e:
            raise ValueError(f"Failed to convert Word document: {str(e)}")

    def preview(self, file_path):
        """Generate preview of Word document"""
        try:
            doc = Document(file_path)
            preview = ""
            
            # Get first few paragraphs
            for paragraph in doc.paragraphs[:5]:  # First 5 paragraphs
                preview += paragraph.text + "\n"
                if len(preview) > 1000:
                    break
                    
            # Add table preview if exists
            if doc.tables and len(preview) < 800:
                preview += "\n[Table Contents]\n"
                table = doc.tables[0]  # Preview first table
                for i, row in enumerate(table.rows[:3]):  # First 3 rows
                    if len(preview) > 950:
                        break
                    preview += " | ".join(cell.text for cell in row.cells) + "\n"
                if len(table.rows) > 3:
                    preview += "...\n"
                    
            return preview[:1000]
            
        except Exception as e:
            raise ValueError(f"Failed to preview Word document: {str(e)}")
